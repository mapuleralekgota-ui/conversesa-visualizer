from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageEnhance


BASE = Path(__file__).resolve().parent
ASSETS = BASE / "real-assets"
OUT = BASE / "Converse_SA_Visualizer.docx"
DOC_ASSETS = BASE / "doc-assets"

BLACK = RGBColor(7, 7, 7)
CREAM = RGBColor(246, 240, 229)
ACID = RGBColor(196, 211, 72)
ORANGE = RGBColor(239, 91, 42)
GRAY = RGBColor(82, 82, 82)


def crop_banner(src_name, out_name, size=(1800, 720), brightness=0.74, contrast=1.08):
    DOC_ASSETS.mkdir(exist_ok=True)
    src = Image.open(ASSETS / src_name).convert("RGB")
    src_ratio = src.width / src.height
    target_ratio = size[0] / size[1]
    if src_ratio > target_ratio:
        new_width = int(src.height * target_ratio)
        left = (src.width - new_width) // 2
        src = src.crop((left, 0, left + new_width, src.height))
    else:
        new_height = int(src.width / target_ratio)
        top = (src.height - new_height) // 2
        src = src.crop((0, top, src.width, top + new_height))
    src = src.resize(size, Image.Resampling.LANCZOS)
    src = ImageEnhance.Brightness(src).enhance(brightness)
    src = ImageEnhance.Contrast(src).enhance(contrast)
    out = DOC_ASSETS / out_name
    src.save(out, quality=92)
    return out


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D9D2C8", size="8"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def run_font(run, size=None, bold=None, color=None, italic=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text="", size=11, color=BLACK, bold=False, italic=False, after=6, before=0, align=None, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, eyebrow, title):
    add_para(doc, eyebrow.upper(), size=8.5, bold=True, color=ORANGE, after=2)
    p = add_para(doc, after=10)
    for idx, part in enumerate(title.split("|")):
        run = p.add_run(part.strip())
        run_font(run, size=22 if idx == 0 else 21, bold=True, color=BLACK if idx == 0 else ACID)
        if idx == 0 and len(title.split("|")) > 1:
            p.add_run(" ")
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12
        run = p.add_run(item)
        run_font(run, size=10.2, color=BLACK)


def add_card_table(doc, cards, columns=2, fill="111111", text_color=CREAM):
    rows = (len(cards) + columns - 1) // columns
    table = doc.add_table(rows=rows, cols=columns)
    table.autofit = False
    set_table_width(table)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_shading(cell, fill)
            set_cell_border(cell, color="2D2D2D")
            set_cell_margins(cell, top=160, bottom=160, start=180, end=180)
    for idx, card in enumerate(cards):
        cell = table.cell(idx // columns, idx % columns)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(card[0].upper())
        run_font(r, size=10, bold=True, color=ACID)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(card[1])
        run_font(r2, size=9.2, color=text_color)
    add_para(doc, after=8)
    return table


def add_image(doc, image_path, width=6.5, after=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    return p


def add_page_break(doc):
    doc.add_page_break()


def build():
    shoes = crop_banner("real-shoes.jpg", "doc-shoes.jpg", size=(1800, 760), brightness=0.68)
    crowd = crop_banner("real-crowd.jpg", "doc-crowd.jpg", size=(1800, 660), brightness=0.68)
    street = crop_banner("real-street.jpg", "doc-street.jpg", size=(1800, 660), brightness=0.74)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    header = section.header.paragraphs[0]
    header.text = "Converse SA Visualizer | Impacta"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_font(header.runs[0], size=8.5, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.text = "Client-safe outline | May 2026"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_font(footer.runs[0], size=8.5, color=GRAY)

    add_image(doc, shoes, width=6.5, after=22)
    add_para(doc, "IMPACTA VISUALIZER", size=9, bold=True, color=ORANGE, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "CONVERSE SA", size=26, bold=True, color=BLACK, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Past. Present. Always You.", size=18, bold=True, color=ACID, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "A culture-first outline for conversation, refinement and partner alignment.", size=11.5, color=GRAY, italic=True, after=24, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_card_table(doc, [
        ("Purpose", "A revised, minimal visualizer that can be shared as a client-safe discussion document."),
        ("Positioning", "Converse as a cultural platform rooted in heritage, self-expression and youth creativity."),
        ("Tone", "Vague by design: enough direction to excite, without locking talent names or final executions."),
        ("Output", "A written summary to support the live visualizer or act as an attachable leave-behind."),
    ], columns=2)

    add_page_break(doc)

    add_heading(doc, "01 / Overview", "What we are | exploring")
    add_para(doc, "A culture-first approach that strengthens Converse's relevance with South African youth by meeting them in the spaces, stories and moments that shape their world.", size=11.2, after=6)
    add_para(doc, "This is not just about campaigns. It is about building cultural connections that last.", size=11.2, bold=True, after=12)
    add_image(doc, street, width=6.5, after=14)

    add_heading(doc, "02 / Our approach", "Not just campaigns. | A cultural ecosystem.")
    add_card_table(doc, [
        ("Community first", "Build with the culture, not just for it."),
        ("Emotional relevance", "Connect Converse to identity, not just product."),
        ("On the ground", "Create real experiences in real communities."),
        ("Content that lives", "Shape storytelling that feels organic, raw and shareable."),
        ("Past. Present. Always.", "Honour heritage while driving the future."),
    ], columns=2)

    add_page_break(doc)

    add_heading(doc, "03 / Our audience", "Broad. Diverse. | Culturally connected.")
    add_para(doc, "The opportunity is to reach students, creatives, trendsetters, music lovers, skaters and everyday self-expressers through spaces where style and identity already live.", size=11.2, after=8)
    add_bullets(doc, [
        "Students and young creatives",
        "Music and culture communities",
        "Campus and street-style audiences",
        "Everyday self-expressers across different identities",
    ])

    add_heading(doc, "04 / Past meets present", "Our story. Our culture. | Our future.")
    add_para(doc, "From the streets to today's creators, Converse has always been part of the story. The opportunity is to bridge heritage with modern youth culture, celebrating the icons who came before us and the voices shaping what is next.", size=11.2, after=12)
    add_image(doc, shoes, width=6.5, after=14)

    add_page_break(doc)

    add_heading(doc, "05 / Brand ambassador idea", "A voice of | this generation.")
    add_para(doc, "A talent or creator who moves culture forward and brings Converse along for the journey. The direction remains intentionally open until the right person is aligned.", size=11.2, after=8)
    add_bullets(doc, [
        "Content takeover",
        "Music integration",
        "Dance or movement-led challenge",
        "Campus activations",
        "Youth participation moments",
    ])

    add_heading(doc, "06 / Experience pillars", "Culture you can | step into.")
    add_card_table(doc, [
        ("Music culture", "Live moments, gigs, sound and discovery."),
        ("Fashion and style", "Self-expression through what people wear."),
        ("Campus culture", "Student energy, activation and real connection."),
        ("Creators and content", "Collaborations that influence and inspire."),
        ("Community activation", "Platforms for youth to participate."),
        ("Digital storytelling", "Stories that travel, live and get shared."),
    ], columns=2)

    add_page_break(doc)

    add_heading(doc, "07 / Experience concepts", "Big moments. | Real impact.")
    add_card_table(doc, [
        ("Creative workshops and labs", "Hands-on spaces for youth creativity."),
        ("Community events", "Local experiences with cultural credibility."),
        ("Converse campus connect", "Campus-led activation moments."),
        ("Local music showcases", "Sound-led moments that feel authentic."),
        ("Pop-up installations", "Physical environments people want to enter."),
        ("Digital content series", "Stories designed to live after the event."),
    ], columns=2)
    add_image(doc, crowd, width=6.5, after=14)

    add_page_break(doc)

    add_heading(doc, "08 / Why this works", "We do not follow culture. | We flow with it.")
    add_para(doc, "The Converse consumer already lives in these spaces. Impacta's role is to show up authentically, add value and create impact that lasts.", size=11.2, after=8)
    add_bullets(doc, [
        "Authentic presence in spaces that already matter",
        "A stronger bridge between heritage and modern youth culture",
        "Content moments that travel naturally",
        "A platform that can grow into future campaigns and partnerships",
    ])

    add_heading(doc, "09 / Impacta's role", "From strategy to | culture in action.")
    add_para(doc, "Impacta shapes the thinking, designs the experience and manages delivery so the idea feels culturally sharp and commercially clear.", size=11.2, after=8)
    add_bullets(doc, [
        "Cultural strategy",
        "Experiential design",
        "Creator partnerships",
        "Production and execution",
        "Content direction",
    ])

    add_page_break(doc)

    add_heading(doc, "10 / Next steps", "Let us build what is next. | Together.")
    add_para(doc, "We would love to unpack these ideas further and explore how we can create something meaningful for Converse South Africa.", size=11.6, after=16)
    add_card_table(doc, [
        ("Conversation", "Use this as a direction-setting discussion document."),
        ("Refinement", "Align around the territories that feel strongest for Converse SA."),
        ("Development", "Build out the preferred concepts into a more formal proposal."),
        ("Execution", "Shape a campaign system that is creative, practical and culturally relevant."),
    ], columns=2)
    add_para(doc, "Prepared by Impacta", size=10, bold=True, color=GRAY, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
